import os
from dotenv import load_dotenv
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma

try:
    from st_aggrid import AgGrid, GridOptionsBuilder  # optionales Grid
except ImportError:
    AgGrid = None

load_dotenv()
#haha nyan

# --- DB laden ---
@st.cache_data(show_spinner=False)
def load_db(persist_dir: str, _embeds: OpenAIEmbeddings):
    vectordb = Chroma(persist_directory=persist_dir, embedding_function=_embeds)
    col = vectordb._collection
    data = col.get(limit=col.count())
    return data['documents'], data['metadatas']


# --- DataFrame für Explorer erzeugen ---
def load_db_entries(persist_dir: str, embeds: OpenAIEmbeddings):
    docs, metas = load_db(persist_dir, embeds)
    rows = []
    for idx, (text, meta) in enumerate(zip(docs, metas)):
        raw_roles = meta.get('roles', '')
        roles_str = ", ".join(raw_roles) if isinstance(raw_roles, list) else raw_roles

        typ = meta.get('type', 'module')
        if typ == "module":
            title = meta.get("module_id", "")
            if meta.get("module_title"):
                title += " – " + meta["module_title"]
        elif typ == "threat":
            title = meta.get('threat_title', '')
        else:
            title = meta.get("requirement_title", meta.get("requirement_id", ""))

        rows.append({
            "ID": idx,
            "type": typ,
            "chapter": meta.get("chapter_id", ""),
            "module": meta.get("module_id", ""),
            "title": title,
            "requirement_id": meta.get("requirement_id", ""),
            "threat_title": meta.get("threat_title", ""),
            "level": meta.get("level", ""),
            "roles": roles_str,
            "snippet": (text.replace("\n", " ")[:300] + "…") if len(text) > 300 else text
        })

    df = pd.DataFrame(rows)
    return df, docs, metas


# --- Hierarchie aus Metadaten bauen ---
@st.cache_data(show_spinner=False)
def build_hierarchy(docs, metas):
    hier = {}
    for idx, (doc, meta) in enumerate(zip(docs, metas)):
        chap = meta.get('chapter_id', 'UNKNOWN')
        mod = meta.get('module_id', 'UNKNOWN')
        typ = meta.get('type', 'module')
        chapter = hier.setdefault(chap, {})
        module = chapter.setdefault(mod, {'module_docs': [], 'requirements': {}})
        if typ == 'module':
            module['module_docs'].append((idx, doc, meta))
        elif typ == 'requirement':
            rid = meta.get('requirement_id', 'UNKNOWN')
            req = module['requirements'].setdefault(rid, {'chunks': [], 'meta': meta})
            req['chunks'].append((idx, doc))
    return hier


def build_docx(requirements):
    """
    Sortiert die Anforderungen nach Kapitel, Modul, ID und erstellt ein DOCX.
    """
    # Sortiere Requirements
    sorted_reqs = sorted(
        requirements,
        key=lambda r: (
            r['meta'].get('chapter_id', ''),
            r['meta'].get('module_id', ''),
            r['meta'].get('requirement_id', '')
        )
    )
    doc = Document()
    doc.add_heading("Ausgewählte Anforderungen", level=1)

    current_chap = None
    for req in sorted_reqs:
        meta = req['meta']
        chap = meta.get('chapter_id')
        # Page break on chapter change (except first)
        if current_chap is not None and chap != current_chap:
            doc.add_page_break()
        current_chap = chap

        title = meta.get('requirement_title', meta.get('requirement_id'))
        doc.add_heading(title, level=2)
        for _, chunk in req['chunks']:
            doc.add_paragraph(chunk)
        doc.add_paragraph(
            f"ID: {meta.get('requirement_id')}  |  Level: {meta.get('level')}  |  Rollen: {meta.get('roles')}",
            style="IntenseQuote"
        )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def module_matches(info, query):
    _, desc, meta = info['module_docs'][0]
    if query in meta.get('module_title', '').lower(): return True
    if query in desc.lower(): return True
    for rid, req in info['requirements'].items():
        title = req['meta'].get('requirement_title', '').lower()
        if query in title: return True
        for _, chunk in req['chunks']:
            if query in chunk.lower(): return True
    return False


def add_to_cart(meta, chunks):
    entry = {'meta': meta, 'chunks': chunks}
    if entry not in st.session_state.cart:
        st.session_state.cart.append(entry)


# --- Streamlit App ---
def main():
    st.set_page_config(page_title='Kompendium Explorer', layout='wide')
    st.title('IT-Grundschutz Kompendium – Explorer')

    # Sidebar: Ausgewählte Anforderungen
    st.sidebar.header("Ausgewählte Anforderungen")
    if 'cart' not in st.session_state:
        st.session_state.cart = []
    for i, item in enumerate(st.session_state.cart):
        title = item['meta'].get('requirement_title', item['meta']['requirement_id'])
        cols = st.sidebar.columns([0.8, 0.2])
        cols[0].write(f"• {title}")
        if cols[1].button("✕", key=f"rem_{i}"):
            st.session_state.cart.pop(i)
    if st.session_state.cart:
        buf = build_docx(st.session_state.cart)
        st.sidebar.download_button("Download DOCX", data=buf, file_name="Anforderungen.docx")
    else:
        st.sidebar.write("_Keine Anforderungen ausgewählt_")

    # Navigation
    page = st.sidebar.radio('Navigation', [
        'Datenbank Explorer',
        'Semantische Suche / Q&A',
        'Drilldown der Bausteine'
    ])

    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        st.error('OPENAI_API_KEY fehlt in der Umgebung')
        return
    embeds = OpenAIEmbeddings(openai_api_key=api_key)

    # Datenbank Explorer
    if page == 'Datenbank Explorer':
        st.header('Datenbank Explorer')
        df, docs, metas = load_db_entries('db', embeds)
        st.markdown(f"**Gesamt:** {len(df)} Dokumente in der Vektor-DB")
        with st.sidebar:
            st.subheader('Filter')
            type_sel = st.multiselect('Dokument-Typ', sorted(df['type'].unique()), sorted(df['type'].unique()))
            chapters_sel = st.multiselect('Kapitel', sorted(df['chapter'].unique()), sorted(df['chapter'].unique()))
            modules_sel = st.multiselect('Baustein-ID', sorted(df['module'].unique()), sorted(df['module'].unique()))
            level_sel = st.multiselect('Level (B/S/H)', sorted(df['level'].unique()), sorted(df['level'].unique()))
            roles_sel = st.multiselect(
                'Rollen', sorted({r for row in df['roles'] for r in row.split(', ') if r}),
                sorted({r for row in df['roles'] for r in row.split(', ') if r})
            )
            sort_col = st.selectbox('Sortiere nach', ['chapter', 'module', 'type', 'level', 'title'], index=0)
            ascending = st.checkbox('Aufsteigend', True)
        mask = (
                df['type'].isin(type_sel) &
                df['chapter'].isin(chapters_sel) &
                df['module'].isin(modules_sel) &
                df['level'].isin(level_sel) &
                ((df['type'] != 'requirement') |
                 df['roles'].apply(lambda rs: any(r in rs for r in roles_sel)))
        )
        df_filt = df[mask].sort_values(by=sort_col, ascending=ascending).reset_index(drop=True)
        st.subheader('Gefilterte Dokumente')
        if AgGrid:
            gb = GridOptionsBuilder.from_dataframe(df_filt)
            gb.configure_default_column(editable=False, wrapText=True, autoHeight=True)
            gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=15)
            grid = AgGrid(df_filt, gridOptions=gb.build(), height=600, fit_columns_on_grid_load=True)
            sel = grid.get('selected_rows') or []
            if sel:
                idx = sel[0]['_selectedRowNodeInfo']['nodeRowIndex']
                st.markdown('---')
                st.subheader('Detail')
                st.write(docs[idx])
        else:
            st.dataframe(df_filt, height=600)

    # Q&A
    elif page == 'Semantische Suche / Q&A':
        st.header('Semantische Suche / Q&A')
        only_req = st.checkbox('Nur nach Requirements suchen', value=False)
        query = st.text_input('Suche / Frage eingeben:')
        k = st.slider('Anzahl Ergebnisse', 1, 20, 5)
        if query:
            vdb = Chroma(persist_directory='db', embedding_function=embeds)
            if only_req:
                results = vdb.max_marginal_relevance_search(query, k=k, fetch_k=k * 5, lambda_mult=0.7,
                                                            filter={"type": "requirement"})
            else:
                results = vdb.max_marginal_relevance_search(query, k=k, fetch_k=k * 5, lambda_mult=0.7)
            for i, doc in enumerate(results, 1):
                meta = doc.metadata
                header = f"**{i}.** {meta.get('module_id', '–')} • {meta.get('type', '–')}"
                if meta.get('requirement_title'):
                    header += f" • {meta['requirement_title']}"
                elif meta.get('requirement_id'):
                    header += f" • {meta['requirement_id']}"
                if meta.get('threat_title'):
                    header += f" • THREAT: {meta['threat_title']}"
                st.markdown(header)
                st.write(doc.page_content)
                st.caption(meta)
                st.markdown('---')

    # Drilldown der Bausteine
    else:
        st.header('Drilldown der Bausteine')
        query = st.text_input("Schnellsuche (z.B. 'Server')").strip().lower()
        docs, metas = load_db('db', embeds)
        hier = build_hierarchy(docs, metas)
        for chap_id, modules in sorted(hier.items()):
            if query and not any(module_matches(info, query) for info in modules.values()):
                continue
            chap_title = next((m.get('chapter_title') for m in metas if m.get('chapter_id') == chap_id), chap_id)
            if not st.checkbox(f"{chap_id} – {chap_title}", key=f'chap_{chap_id}', value=bool(query)):
                continue
            # module level
            for mod_id, info in sorted(modules.items()):
                if query and not module_matches(info, query): continue
                _, desc, mod_meta = info['module_docs'][0]
                label = f"{mod_id} – {mod_meta.get('module_title', mod_id)}"
                with st.expander(label, expanded=bool(query)):
                    st.write(desc)
                    st.markdown("**Anforderungen:**")
                    for rid, req in sorted(info['requirements'].items()):
                        meta = req['meta']
                        title = meta.get('requirement_title', rid)
                        is_match = not query or (
                                    query in title.lower() or any(query in c.lower() for _, c in req['chunks']))
                        if not is_match:
                            continue
                        # draw title and add-button, use checkbox for toggle when not searching
                        cols = st.columns([0.7, 0.15, 0.15])
                        if query:
                            cols[0].markdown(f"- **{title}**")
                            if cols[1].button("＋", key=f"add_{chap_id}_{mod_id}_{rid}", on_click=add_to_cart,
                                              args=(meta, req['chunks'])):
                                pass
                            # show content immediately
                            for _, chunk in req['chunks']:
                                st.write(chunk)
                            st.caption(meta)
                        else:
                            checked = cols[0].checkbox(title, key=f"tog_{chap_id}_{mod_id}_{rid}")
                            if cols[1].button("＋", key=f"add_{chap_id}_{mod_id}_{rid}", on_click=add_to_cart,
                                              args=(meta, req['chunks'])):
                                pass
                            if checked:
                                for _, chunk in req['chunks']:
                                    st.write(chunk)
                                st.caption(meta)


if __name__ == '__main__':
    main()
